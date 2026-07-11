# -*- coding: utf-8 -*-
"""
Passo 7 — Relatório final .docx no formato do REPORT_TIPO.

Requer:
  results/*.csv, results/stability.json, results/figures/*.png
  results/topic_labels.json      {"0": "label", ...}
  results/report_content.json    (textos interpretativos)

Output: LDA_Report_JoaoLuis.docx (na raiz do repositório)
"""
import json
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt

BASE = Path(__file__).resolve().parents[1]   # raiz do repositório
RES = BASE / "results"
FIG = RES / "figures"
OUT = BASE / "LDA_Report_JoaoLuis.docx"

# ---------------------------------------------------------------- dados
stab = json.loads((RES / "stability.json").read_text())
K = stab["k"]
labels = {int(k): v for k, v in json.loads(
    (RES / "topic_labels.json").read_text(encoding="utf-8")).items()}
C = json.loads((RES / "report_content.json").read_text(encoding="utf-8"))

meta = pd.read_csv(BASE / "data" / "metadata.csv")
sweep = pd.read_csv(RES / "coherence_sweep.csv")
overview = pd.read_csv(RES / "topic_overview.csv")
terms = pd.read_csv(RES / "topic_terms.csv")
dom = pd.read_csv(RES / "dominant_topics.csv")
dt = pd.read_csv(RES / "doc_topic_matrix.csv")

N_DOCS = len(meta)
Y0, Y1 = int(meta.year.min()), int(meta.year.max())
KMAX_CV = int(sweep.loc[sweep.cv_mean.idxmax(), "k"])


def clean_authors(a: str) -> str:
    return str(a).replace(" e ", " and ").strip()


def top_terms(t, n=12):
    tt = terms[(terms.topic == t) & (terms["rank"] <= n)].sort_values("rank")
    return ", ".join(tt.term)


# ---------------------------------------------------------------- helpers
doc = Document()
doc.core_properties.author = "João Luís"
doc.core_properties.title = "Topic Modeling (LDA) — Port Digitalization Corpus"

normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(11)
pf = normal.paragraph_format
pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
pf.line_spacing = 1.15
pf.space_after = Pt(10)


def para(text, style=None, bold_lead=None):
    p = doc.add_paragraph(style=style)
    if bold_lead:
        r = p.add_run(bold_lead + " ")
        r.bold = True
    p.add_run(text)
    return p


def bullet(text, bold_lead=None):
    return para(text, style="List Bullet", bold_lead=bold_lead)


def caption(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(9)
    return p


def picture(path, width=6.2):
    doc.add_picture(str(path), width=Inches(width))


def add_toc_field():
    p = doc.add_paragraph()
    run = p.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t")
    t.text = "Right-click here and choose 'Update Field' to build the table of contents."
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    for el in (fld_begin, instr, fld_sep, t, fld_end):
        run._r.append(el)


def table(headers, rows, widths=None, bold_cells=None, font_size=None):
    """bold_cells: set de (row_idx, col_idx) a negrito (0-based, sem header)."""
    tb = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tb.style = "Table Grid"
    for j, h in enumerate(headers):
        cell = tb.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        r.bold = True
        if font_size:
            r.font.size = Pt(font_size)
        cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    for i, row in enumerate(rows):
        for j, v in enumerate(row):
            cell = tb.rows[i + 1].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(v))
            if bold_cells and (i, j) in bold_cells:
                r.bold = True
            if font_size:
                r.font.size = Pt(font_size)
            cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    if widths:
        for j, w in enumerate(widths):
            for row in tb.rows:
                row.cells[j].width = Inches(w)
    return tb


# ---------------------------------------------------------------- título
p = doc.add_paragraph()
r = p.add_run("TOPIC MODELING (LDA)")
r.bold = True
r.font.size = Pt(20)
para(C["subtitle"])
para(C["description"])
para(f"Corpus: {N_DOCS} articles · full text · {Y0}–{Y1}")
para(f"Result: {K} latent topics (k chosen by C_V coherence and interpretability)")
para(f"Prepared for: {C['prepared_for']} · {C['date_line']}")

# ---------------------------------------------------------------- TOC
doc.add_heading("Contents", level=1)
add_toc_field()

# ---------------------------------------------------------------- exec summary
doc.add_heading("Executive summary", level=1)
para(C["exec_intro"])
order = overview.sort_values("prevalence", ascending=False).topic.tolist()
for t in sorted(labels):
    lab = labels[t]
    if len(lab) > 1 and lab[1].islower():  # não descapitalizar acrónimos
        lab = lab[0].lower() + lab[1:]
    bullet(f"T{t} — {lab}.")
para("Key findings:")
for kf in C["key_findings"]:
    bullet(kf["text"], bold_lead=kf["lead"])

# ---------------------------------------------------------------- 1-2
doc.add_heading("1. Introduction and objective", level=1)
para(C["introduction"])

doc.add_heading("2. Corpus and data", level=1)
para(C["corpus_paragraph"])
table(
    ["Indicator", "Value"],
    [
        ["Number of articles", str(N_DOCS)],
        ["Publication period", f"{Y0}–{Y1}"],
        ["Dictionary terms (after filtering)", f"{stab['n_terms']:,}"],
        ["Inference engine", "gensim LdaModel (Variational Bayes)"],
        ["Priors (alpha / beta)", f"{stab['alpha']} / {stab['eta']}"],
        ["Number of topics (k)", str(K)],
    ],
    widths=[3.4, 2.8],
)

# ---------------------------------------------------------------- 3
doc.add_heading("3. Methodology", level=1)
doc.add_heading("3.1 Preprocessing", level=2)
for b in C["preprocessing_bullets"]:
    bullet(b["text"], bold_lead=b["lead"])
doc.add_heading("3.2 Dictionary and bag-of-words representation", level=2)
para(C["dictionary_paragraph"])
doc.add_heading("3.3 Inference and hyperparameters", level=2)
para(C["inference_paragraph"])
doc.add_heading("3.4 Selecting the number of topics", level=2)
para(C["k_selection_paragraph"])

# ---------------------------------------------------------------- 4
doc.add_heading("4. Selection of the number of topics", level=1)
para(C["selection_justification"])
picture(FIG / "fig1_coherence.png")
caption("Figure 1 — C_V coherence vs. number of topics (mean ± sd over 3 initializations).")
rows = []
for _, r_ in sweep.iterrows():
    kk = int(r_.k)
    mark = f"{kk}  ◄ chosen" if kk == K else str(kk)
    rows.append([mark, f"{r_.cv_mean:.4f}", f"{r_.cv_std:.4f}"])
table(["k (topics)", "Mean C_V", "Std. dev."], rows, widths=[2.0, 2.0, 2.0])

# ---------------------------------------------------------------- 5
doc.add_heading(f"5. Results — the {C['k_word']} topics", level=1)
para(C["results_explainer"])
picture(FIG / "fig2_wordclouds.png")
caption(f"Figure 2 — Dominant terms of each of the {C['k_word']} topics.")
rows = []
for t in order:
    ov = overview[overview.topic == t].iloc[0]
    rows.append([f"T{t}", labels[t], int(ov.n_dominant), f"{ov.prevalence*100:.1f}%"])
table(["#", "Topic", "Articles", "Prevalence"], rows, widths=[0.6, 3.6, 1.0, 1.2])

for t in order:
    ov = overview[overview.topic == t].iloc[0]
    doc.add_heading(f"Topic T{t} — {labels[t]}", level=2)
    para(f"Top terms: {top_terms(t)}")
    para(f"Size: {int(ov.n_dominant)} dominant articles · "
         f"Prevalence: {ov.prevalence*100:.1f}%")
    para(C["topics"][str(t)]["interpretation"])
    para("Representative articles:")
    reps = dom[dom.dominant_topic == t].sort_values(
        "proportion", ascending=False).head(4)
    for _, r_ in reps.iterrows():
        bullet(f"{clean_authors(r_.authors)} ({int(r_.year)}) — {r_.title}  "
               f"(proportion {r_.proportion:.2f})")

# ---------------------------------------------------------------- 6
doc.add_heading("6. Distribution and temporal evolution", level=1)
picture(FIG / "fig3_prevalence.png")
caption("Figure 3 — Topic prevalence across the corpus.")
picture(FIG / "fig4_heatmap.png")
caption("Figure 4 — Heatmap of the mean share of each topic per publication year.")
para(C["temporal_interpretation"])
bullet(C["temporal_note"], bold_lead="Note.")
picture(FIG / "fig5_papers_year.png")
caption("Figure 5 — Number of articles per publication year.")

# ---------------------------------------------------------------- 7
doc.add_heading("7. Geographic focus", level=1)
picture(FIG / "fig6_countries.png")
caption("Figure 6 — Most-mentioned countries in the article bodies (top 15).")
para(C["geographic_paragraph"])

# ---------------------------------------------------------------- 8
doc.add_heading("8. Robustness and limitations", level=1)
para("Robustness.")
for b in C["robustness_bullets"]:
    bullet(b["text"], bold_lead=b["lead"])
para("Limitations (methodological transparency).")
for b in C["limitations_bullets"]:
    bullet(b["text"], bold_lead=b["lead"])

# ---------------------------------------------------------------- 9
doc.add_heading("9. Conclusions", level=1)
para(C["conclusions"])

# ---------------------------------------------------------------- Apêndice A
doc.add_heading("Appendix A — Dominant topic of each article", level=1)
para(C["appendixA_explainer"])
adf = dom.copy()
adf = adf.sort_values(["dominant_topic", "proportion"],
                      ascending=[True, False]).reset_index(drop=True)
rows = []
for i, r_ in adf.iterrows():
    rows.append([
        str(i + 1),
        f"{clean_authors(r_.authors)} ({int(r_.year)})",
        str(r_.title)[:90],
        f"T{int(r_.dominant_topic)}",
        f"{r_.proportion*100:.0f}%",
    ])
table(["#", "Authors (year)", "Title", "Top.", "Prop."],
      rows, widths=[0.4, 1.7, 3.0, 0.5, 0.6])

# ---------------------------------------------------------------- Apêndice B
# secção horizontal: 12 colunas de probabilidades não cabem em retrato
from docx.enum.section import WD_ORIENT, WD_SECTION_START

sec = doc.add_section(WD_SECTION_START.NEW_PAGE)
sec.orientation = WD_ORIENT.LANDSCAPE
sec.page_width, sec.page_height = sec.page_height, sec.page_width

doc.add_heading("Appendix B — Document–topic distribution matrix", level=1)
para(C["appendixB_explainer"])
bdf = dt.copy()
bdf["study"] = bdf.apply(
    lambda r_: f"({clean_authors(r_.authors)}, {int(r_.year)})", axis=1)
bdf = bdf.sort_values("study").reset_index(drop=True)
tcols = [f"T{t}" for t in range(K)]
rows, bold_cells = [], set()
for i, r_ in bdf.iterrows():
    vals = [r_[c] for c in tcols]
    dominant = max(range(K), key=lambda j: vals[j])
    bold_cells.add((i, 1 + dominant))
    rows.append([r_.study[:38]] + [f"{v:.6f}" for v in vals])
table(["Study"] + tcols, rows,
      widths=[1.9] + [0.58] * K, bold_cells=bold_cells, font_size=8)

# corrige o template do python-docx: <w:zoom> requer o atributo percent
zoom = doc.settings.element.find(qn("w:zoom"))
if zoom is not None and zoom.get(qn("w:percent")) is None:
    zoom.set(qn("w:percent"), "100")

doc.save(OUT)
print("Relatório gravado em", OUT)
